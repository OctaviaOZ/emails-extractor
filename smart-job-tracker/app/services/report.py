import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from datetime import datetime
import os
from app.models import ApplicationStatus
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_status_label(status, mapping=None):
    """
    Returns the localized label for a status using the provided mapping.
    Defaults to the English enum value if no mapping is found.
    """
    if mapping:
        # Check for direct match (e.g., 'Applied')
        if status.value in mapping:
            return mapping[status.value]
        # Check for match by key (e.g., if mapping uses 'Applied' as key)
        if status in mapping:
            return mapping[status]
            
    # Default Fallback (English / Raw Value)
    return status.value

def filter_applications_by_date(applications, start_date=None, end_date=None):
    """
    Filters applications based on last_updated date.
    """
    if not start_date and not end_date:
        return applications
    
    filtered = []
    start = pd.to_datetime(start_date) if start_date else datetime.min
    end = pd.to_datetime(end_date) if end_date else datetime.max
    
    # Ensure start/end are timezone-naive or aware matching the app dates
    # Assuming app.last_updated is naive UTC as per models.py
    if start.tzinfo: start = start.replace(tzinfo=None)
    if end.tzinfo: end = end.replace(tzinfo=None)
    
    # Make end date inclusive (end of day)
    end = end.replace(hour=23, minute=59, second=59)

    for app in applications:
        # Check if last_updated falls in range
        if start <= app.last_updated <= end:
            filtered.append(app)
    return filtered

def generate_word_report(applications, output_filename, start_date=None, end_date=None, config=None):
    """
    Generates a MS Word report (.docx) in German with date filtering.
    """
    # Filter Data
    filtered_apps = filter_applications_by_date(applications, start_date, end_date)
    
    if not filtered_apps:
        return None

    # Convert to DataFrame
    data = [app.model_dump() for app in filtered_apps]
    df = pd.DataFrame(data)
    
    if df.empty:
        return None

    df['last_updated'] = pd.to_datetime(df['last_updated'])
    df = df.sort_values(by='last_updated', ascending=False)

    # German Mapping
    german_mapping = {
        "APPLIED": "Beworben",
        "INTERVIEW": "Interview",
        "ASSESSMENT": "Assessment",
        "PENDING": "Laufend",
        "OFFER": "Angebot",
        "REJECTED": "Abgelehnt",
        "UNKNOWN": "Unbekannt"
    }

    # Prepare Report Data
    report_data = []
    status_counts = {k: 0 for k in german_mapping.values()}

    for company, group in df.groupby('company_name'):
        latest = group.iloc[0]
        interview_count = group[group['status'] == ApplicationStatus.INTERVIEW].shape[0]
        status_label = get_status_label(latest['status'], german_mapping)
        
        # Increment counter
        if status_label in status_counts:
            status_counts[status_label] += 1
        else:
            status_counts[status_label] = 1

        report_data.append({
            "Firma": company,
            "Status": status_label,
            "Interviews": interview_count,
            "Letzter Kontakt": latest['last_updated'].strftime('%d.%m.%Y'),
            "_sort_date": latest['last_updated']
        })

    report_df = pd.DataFrame(report_data).sort_values(by='_sort_date', ascending=False)

    # Create Document
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Bewerbungsbericht", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date Range Subtitle
    date_str = datetime.now().strftime('%d.%m.%Y')
    if start_date and end_date:
        date_str = f"{start_date.strftime('%d.%m.%Y')} - {end_date.strftime('%d.%m.%Y')}"
    elif start_date:
        date_str = f"Ab {start_date.strftime('%d.%m.%Y')}"
    
    p = doc.add_paragraph(f"Zeitraum: {date_str}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph() # Spacer

    # Main Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Header
    hdr_cells = table.rows[0].cells
    headers = ['Firma', 'Status', 'Interviews', 'Letzter Kontakt']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True

    # Rows
    for _, row in report_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Firma'])
        row_cells[1].text = str(row['Status'])
        row_cells[2].text = str(row['Interviews'])
        row_cells[3].text = str(row['Letzter Kontakt'])

    doc.add_page_break()

    # Summary Section
    doc.add_heading('Zusammenfassung der Status', level=1)
    
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Shading Accent 1'
    
    sum_hdr = summary_table.rows[0].cells
    sum_hdr[0].text = "Status"
    sum_hdr[1].text = "Anzahl"
    sum_hdr[0].paragraphs[0].runs[0].font.bold = True
    sum_hdr[1].paragraphs[0].runs[0].font.bold = True
    
    # Sort summary by count desc
    sorted_counts = sorted(status_counts.items(), key=lambda item: item[1], reverse=True)
    
    total = 0
    for status, count in sorted_counts:
        if count > 0:
            row = summary_table.add_row().cells
            row[0].text = status
            row[1].text = str(count)
            total += count
            
    # Total Row
    total_row = summary_table.add_row().cells
    total_row[0].text = "Gesamt"
    total_row[1].text = str(total)
    total_row[0].paragraphs[0].runs[0].font.bold = True
    total_row[1].paragraphs[0].runs[0].font.bold = True

    doc.save(output_filename)
    return output_filename

def generate_pdf_report(applications, output_filename, config=None):
    """
    Generates a PDF report from a list of JobApplication objects.
    """
    if not applications:
        return None

    # Convert list of JobApplication objects to DataFrame
    data = [app.model_dump() for app in applications]
    df = pd.DataFrame(data)

    if df.empty:
        return None

    # Ensure last_updated is datetime
    df['last_updated'] = pd.to_datetime(df['last_updated'])

    # Sort by date descending to prioritize latest entries
    df = df.sort_values(by='last_updated', ascending=False)

    report_data = []
    
    mapping = config.get('report_mapping', {}) if config else {}

    # Group by company to aggregate data
    # Note: With multi-process support, a company might appear multiple times.
    # This report aggregates by company, showing the latest status.
    for company, group in df.groupby('company_name'):
        # Latest entry (first in sorted group) determines the current status and last contact
        latest = group.iloc[0]
        
        # Calculate Interview Count
        # Count number of applications where status is specifically INTERVIEW
        # Note: This is a simplified metric. Ideally we would query ApplicationEvent.
        interview_count = group[group['status'] == ApplicationStatus.INTERVIEW].shape[0]
        
        status_label = get_status_label(latest['status'], mapping)
        
        report_data.append({
            "Company": company,
            "Status": status_label,
            "Interviews": interview_count,
            "Last Contact": latest['last_updated'].strftime('%d.%m.%Y'),
            "_sort_date": latest['last_updated'] # Helper for sorting
        })

    # Create DataFrame for Report
    report_df = pd.DataFrame(report_data)
    
    # Sort by Last Contact desc
    report_df = report_df.sort_values(by='_sort_date', ascending=False)

    # Generate PDF
    doc = SimpleDocTemplate(output_filename, pagesize=letter)
    elements = []
    
    styles = getSampleStyleSheet()
    title = Paragraph(f"Application Report - {datetime.now().strftime('%Y-%m-%d')}", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))

    # Table Data
    # Headers - customizable via config could be an enhancement, hardcoded English for now
    headers = ['Company', 'Status', 'Interviews', 'Last Contact']
    table_data = [headers]
    
    # Rows
    for _, row in report_df.iterrows():
        table_data.append([
            row['Company'], 
            row['Status'], 
            str(row['Interviews']), 
            row['Last Contact']
        ])

    # Table Style
    table = Table(table_data)
    style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ])
    table.setStyle(style)
    elements.append(table)

    try:
        doc.build(elements)
        return output_filename
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None